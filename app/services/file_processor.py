"""
File Processing Service — extracts content from uploaded files for AI analysis.
Supports: PDF, images, Excel/CSV, Word docs, SQL, and code files.
"""
import io
import csv
import base64
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Supported file types
TEXT_EXTENSIONS = {".txt", ".sql", ".py", ".js", ".ts", ".json", ".xml", ".yaml", ".yml", ".md", ".html", ".css", ".sh", ".bat", ".log", ".env", ".cfg", ".ini", ".toml"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"}
MAX_TEXT_LENGTH = 50000  # Max chars to send to AI


async def extract_text_from_file(file_bytes: bytes, filename: str) -> dict:
    """
    Extract text content from a file for AI analysis.
    Returns: {"type": "text"|"image", "content": str, "filename": str, "summary": str}
    """
    ext = Path(filename).suffix.lower()
    result = {"filename": filename, "type": "text", "content": "", "summary": ""}

    try:
        if ext == ".pdf":
            result["content"] = _extract_pdf(file_bytes)
            result["summary"] = f"PDF document: {filename} ({len(result['content'])} characters)"

        elif ext in IMAGE_EXTENSIONS:
            result["type"] = "image"
            result["content"] = base64.b64encode(file_bytes).decode("utf-8")
            result["summary"] = f"Image: {filename} ({len(file_bytes)} bytes)"
            # Detect image media type
            if ext in {".jpg", ".jpeg"}:
                result["media_type"] = "image/jpeg"
            elif ext == ".png":
                result["media_type"] = "image/png"
            elif ext == ".gif":
                result["media_type"] = "image/gif"
            elif ext == ".webp":
                result["media_type"] = "image/webp"
            else:
                result["media_type"] = "image/png"

        elif ext in {".xlsx", ".xls"}:
            result["content"] = _extract_excel(file_bytes)
            result["summary"] = f"Excel spreadsheet: {filename}"

        elif ext == ".csv":
            result["content"] = _extract_csv(file_bytes)
            result["summary"] = f"CSV file: {filename}"

        elif ext in {".docx", ".doc"}:
            result["content"] = _extract_docx(file_bytes)
            result["summary"] = f"Word document: {filename}"

        elif ext in TEXT_EXTENSIONS:
            result["content"] = file_bytes.decode("utf-8", errors="replace")[:MAX_TEXT_LENGTH]
            result["summary"] = f"Text file: {filename} ({ext})"

        else:
            # Try as text
            try:
                result["content"] = file_bytes.decode("utf-8", errors="replace")[:MAX_TEXT_LENGTH]
                result["summary"] = f"File: {filename}"
            except Exception:
                result["content"] = f"[Binary file: {filename}, {len(file_bytes)} bytes — cannot extract text]"
                result["summary"] = f"Binary file: {filename}"

    except Exception as e:
        logger.error(f"File extraction error for {filename}: {e}")
        result["content"] = f"[Error reading {filename}: {str(e)}]"
        result["summary"] = f"Error reading: {filename}"

    # Truncate if too long
    if result["type"] == "text" and len(result["content"]) > MAX_TEXT_LENGTH:
        result["content"] = result["content"][:MAX_TEXT_LENGTH] + f"\n\n... [Truncated — {filename} is very large]"

    return result


def _extract_pdf(data: bytes) -> str:
    """Extract text from PDF."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(data))
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")
        return "\n\n".join(text_parts) if text_parts else "[PDF has no extractable text — may be scanned/image-based]"
    except ImportError:
        return "[PyPDF2 not installed — cannot read PDF]"
    except Exception as e:
        return f"[PDF read error: {e}]"


def _extract_excel(data: bytes) -> str:
    """Extract data from Excel file as formatted text."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            parts.append(f"=== Sheet: {sheet_name} ({len(rows)} rows) ===")
            for i, row in enumerate(rows[:200]):  # Limit to 200 rows
                cells = [str(c) if c is not None else "" for c in row]
                parts.append(" | ".join(cells))
            if len(rows) > 200:
                parts.append(f"... [{len(rows) - 200} more rows]")
        wb.close()
        return "\n".join(parts) if parts else "[Empty spreadsheet]"
    except ImportError:
        return "[openpyxl not installed — cannot read Excel]"
    except Exception as e:
        return f"[Excel read error: {e}]"


def _extract_csv(data: bytes) -> str:
    """Extract data from CSV file."""
    try:
        text = data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = []
        for i, row in enumerate(reader):
            if i >= 200:
                rows.append(f"... [more rows truncated]")
                break
            rows.append(" | ".join(row))
        return "\n".join(rows) if rows else "[Empty CSV]"
    except Exception as e:
        return f"[CSV read error: {e}]"


def _extract_docx(data: bytes) -> str:
    """Extract text from Word document."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs) if paragraphs else "[Empty document]"
    except ImportError:
        return "[python-docx not installed — cannot read Word files]"
    except Exception as e:
        return f"[Word read error: {e}]"


def get_supported_extensions() -> list[str]:
    """Return list of all supported file extensions."""
    return sorted(list(TEXT_EXTENSIONS | IMAGE_EXTENSIONS | {".pdf", ".xlsx", ".xls", ".csv", ".docx", ".doc"}))
